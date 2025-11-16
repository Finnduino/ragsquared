"""Review UI blueprint for viewing audit findings."""

from __future__ import annotations

import json
import io
import os
import tempfile
from datetime import datetime
from pathlib import Path

from flask import Blueprint, jsonify, render_template, request, send_file, Response

from ..config.settings import AppConfig
from ..db.models import Audit, AuditChunkResult, AuditorQuestion, Citation, Flag, Document
from ..db.session import get_session
from ..logging_config import get_logger
from ..services.compliance_score import get_flag_summary
from ..services.final_report_generator import FinalReportGenerator

review_blueprint = Blueprint("review", __name__, url_prefix="/review")
logger = get_logger(__name__)


def _resolve_audit(session, identifier: str) -> Audit | None:
    """Resolve audit by ID or external_id."""
    if identifier.isdigit():
        return session.get(Audit, int(identifier))
    from sqlalchemy import select

    stmt = select(Audit).where(Audit.external_id == identifier)
    return session.execute(stmt).scalar_one_or_none()


# DISABLED: This Flask HTML route conflicts with Next.js frontend
# Next.js handles /review/[auditId] pages via /frontend/app/review/[auditId]/page.tsx
# If you need Flask HTML for CLI/testing, uncomment and access directly via backend URL
# @review_blueprint.route("/<audit_id>")
# def review_audit(audit_id: str):
#     """Render review page for an audit."""
#     session = get_session()
#     audit = _resolve_audit(session, audit_id)
#     
#     if audit is None:
#         return render_template("error.html", message=f"Audit '{audit_id}' not found."), 404
#     
#     # Get filters from query parameters
#     severity_filter = request.args.get("severity", "").upper()
#     regulation_filter = request.args.get("regulation", "")
#     
#     # Build query
#     from sqlalchemy import select
#     
#     query = select(Flag).where(Flag.audit_id == audit.id)
#     
#     if severity_filter:
#         query = query.where(Flag.flag_type == severity_filter)
#     
#     if regulation_filter:
#         query = query.join(Citation).where(
#             Citation.citation_type == "regulation",
#             Citation.reference.ilike(f"%{regulation_filter}%"),
#         )
#     
#     flags = session.execute(query.order_by(Flag.severity_score.desc())).scalars().unique().all()
#     
#     # Get flag summary
#     flag_summary = get_flag_summary(list(flags))
#     
#     # Get auditor questions
#     questions = session.execute(
#         select(AuditorQuestion)
#         .where(AuditorQuestion.audit_id == audit.id)
#         .order_by(AuditorQuestion.priority.asc())
#     ).scalars().all()
#     
#     # Get citations for each flag
#     flag_citations = {}
#     # Get chunk results (which contain context summaries) for each flag
#     flag_contexts = {}
#     for flag in flags:
#         citations = session.execute(
#             select(Citation).where(Citation.flag_id == flag.id)
#         ).scalars().all()
#         flag_citations[flag.id] = citations
#         
#         # Get the chunk result to access context summary
#         chunk_result = session.execute(
#             select(AuditChunkResult).where(
#                 AuditChunkResult.audit_id == audit.id,
#                 AuditChunkResult.chunk_id == flag.chunk_id
#             )
#         ).scalar_one_or_none()
#         
#         if chunk_result and chunk_result.analysis:
#             # Extract context summary from analysis
#             context_summary = chunk_result.analysis.get("context_summary")
#             if context_summary:
#                 flag_contexts[flag.id] = context_summary
#     
#     # Get unique regulation references for filter dropdown
#     all_regulations = session.execute(
#         select(Citation.reference)
#         .join(Flag)
#         .where(Flag.audit_id == audit.id, Citation.citation_type == "regulation")
#         .distinct()
#     ).scalars().all()
#     
#     return render_template(
#         "review.html",
#         audit=audit,
#         flags=flags,
#         flag_summary=flag_summary,
#         questions=questions,
#         flag_citations=flag_citations,
#         flag_contexts=flag_contexts,
#         severity_filter=severity_filter,
#         regulation_filter=regulation_filter,
#         regulations=sorted(set(all_regulations)),
#     )


@review_blueprint.route("/<audit_id>/final-report", methods=["POST"])
def generate_final_report(audit_id: str):
    """Generate a comprehensive final report addressing all compliance issues."""
    session = get_session()
    audit = _resolve_audit(session, audit_id)
    
    if audit is None:
        return jsonify({"error": f"Audit '{audit_id}' not found."}), 404
    
    if audit.status != "completed":
        return jsonify({"error": "Audit must be completed before generating final report."}), 400
    
    try:
        config = AppConfig()
        generator = FinalReportGenerator(session, config)
        report = generator.generate_report(audit.id)
        
        # Convert to dict for JSON response
        return jsonify({
            "audit_id": audit.id,
            "external_id": audit.external_id,
            "executive_summary": report.executive_summary,
            "critical_issues": report.critical_issues,
            "warnings": report.warnings,
            "recommendations": report.recommendations,
            "overall_assessment": report.overall_assessment,
            "raw_content": report.raw_content,
        })
    except Exception as e:
        logger.exception(f"Error generating final report for audit {audit_id}: {e}")
        return jsonify({"error": str(e)}), 500


@review_blueprint.route("/<audit_id>/final-report.json", methods=["GET"])
def download_final_report_json(audit_id: str):
    """Download final report as JSON file."""
    session = get_session()
    audit = _resolve_audit(session, audit_id)
    
    if audit is None:
        return jsonify({"error": f"Audit '{audit_id}' not found."}), 404
    
    if audit.status != "completed":
        return jsonify({"error": "Audit must be completed before generating final report."}), 400
    
    try:
        config = AppConfig()
        generator = FinalReportGenerator(session, config)
        report = generator.generate_report(audit.id)
        
        # Get document name for filename
        document = session.get(Document, audit.document_id) if audit.document_id else None
        audit_name = document.original_filename.replace(Path(document.original_filename).suffix, "") if document else f"audit_{audit.external_id}"
        date_str = datetime.utcnow().strftime("%Y-%m-%d")
        filename = f"final_report_{audit_name}_{date_str}.json"
        
        # Convert to dict
        report_dict = {
            "audit_id": audit.id,
            "external_id": audit.external_id,
            "executive_summary": report.executive_summary,
            "critical_issues": report.critical_issues,
            "warnings": report.warnings,
            "recommendations": report.recommendations,
            "overall_assessment": report.overall_assessment,
            "raw_content": report.raw_content,
            "generated_at": datetime.utcnow().isoformat(),
        }
        
        # Create JSON response with download headers
        json_str = json.dumps(report_dict, indent=2, ensure_ascii=False)
        return Response(
            json_str,
            mimetype="application/json",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        logger.exception(f"Error generating JSON report for audit {audit_id}: {e}")
        return jsonify({"error": str(e)}), 500


@review_blueprint.route("/<audit_id>/final-report.pdf", methods=["GET"])
def download_final_report_pdf(audit_id: str):
    """Download final report as PDF file."""
    session = get_session()
    audit = _resolve_audit(session, audit_id)
    
    if audit is None:
        return jsonify({"error": f"Audit '{audit_id}' not found."}), 404
    
    if audit.status != "completed":
        return jsonify({"error": "Audit must be completed before generating final report."}), 400
    
    try:
        config = AppConfig()
        generator = FinalReportGenerator(session, config)
        report = generator.generate_report(audit.id)
        
        # Convert final report to markdown
        md_content = _final_report_to_markdown(report, audit)
        
        # Generate PDF from markdown
        try:
            from md2pdf.core import md2pdf
        except ImportError:
            return jsonify({"error": "PDF generation requires md2pdf. Install with `pip install md2pdf`."}), 500
        
        # md2pdf needs a file path, so we'll use a temporary approach
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp_md:
            tmp_md.write(md_content)
            tmp_md_path = tmp_md.name
        
        try:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
                tmp_pdf_path = tmp_pdf.name
            
            md2pdf(tmp_pdf_path, source_file=tmp_md_path)
            
            # Read PDF into buffer
            pdf_buffer = io.BytesIO()
            with open(tmp_pdf_path, 'rb') as f:
                pdf_buffer.write(f.read())
            
            pdf_buffer.seek(0)
            
            # Get document name for filename
            document = session.get(Document, audit.document_id) if audit.document_id else None
            audit_name = document.original_filename.replace(Path(document.original_filename).suffix, "") if document else f"audit_{audit.external_id}"
            date_str = datetime.utcnow().strftime("%Y-%m-%d")
            filename = f"final_report_{audit_name}_{date_str}.pdf"
            
            # Clean up temp files
            try:
                os.unlink(tmp_md_path)
                os.unlink(tmp_pdf_path)
            except:
                pass
            
            return send_file(
                pdf_buffer,
                mimetype="application/pdf",
                as_attachment=True,
                download_name=filename
            )
        except Exception as e:
            # Clean up temp files on error
            try:
                os.unlink(tmp_md_path)
                if 'tmp_pdf_path' in locals():
                    os.unlink(tmp_pdf_path)
            except:
                pass
            raise e
            
    except Exception as e:
        logger.exception(f"Error generating PDF report for audit {audit_id}: {e}")
        return jsonify({"error": str(e)}), 500


@review_blueprint.route("/<audit_id>/final-report.docx", methods=["GET"])
def download_final_report_docx(audit_id: str):
    """Download final report as Word DOCX file."""
    session = get_session()
    audit = _resolve_audit(session, audit_id)
    
    if audit is None:
        return jsonify({"error": f"Audit '{audit_id}' not found."}), 404
    
    if audit.status != "completed":
        return jsonify({"error": "Audit must be completed before generating final report."}), 400
    
    try:
        config = AppConfig()
        generator = FinalReportGenerator(session, config)
        report = generator.generate_report(audit.id)
        
        # Generate DOCX using python-docx
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return jsonify({"error": "DOCX generation requires python-docx. Install with `pip install python-docx`."}), 500
        
        doc = DocxDocument()
        
        # Title
        title = doc.add_heading(f"Final Compliance Report: {audit.external_id}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(report.executive_summary)
        
        # Critical Issues
        if report.critical_issues:
            doc.add_heading("Critical Issues", 1)
            for issue in report.critical_issues:
                doc.add_heading(issue.get("title", "Critical Issue"), 2)
                doc.add_paragraph(issue.get("description", ""))
                if issue.get("recommendations"):
                    doc.add_paragraph("Recommendations:")
                    for rec in issue["recommendations"]:
                        doc.add_paragraph(rec, style="List Bullet")
        
        # Warnings
        if report.warnings:
            doc.add_heading("Warnings", 1)
            for warning in report.warnings:
                doc.add_heading(warning.get("title", "Warning"), 2)
                doc.add_paragraph(warning.get("description", ""))
                if warning.get("recommendations"):
                    doc.add_paragraph("Recommendations:")
                    for rec in warning["recommendations"]:
                        doc.add_paragraph(rec, style="List Bullet")
        
        # Recommendations
        if report.recommendations:
            doc.add_heading("Overall Recommendations", 1)
            for rec in report.recommendations:
                doc.add_paragraph(rec, style="List Bullet")
        
        # Overall Assessment
        doc.add_heading("Overall Assessment", 1)
        doc.add_paragraph(report.overall_assessment)
        
        # Save to buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        # Get document name for filename
        document = session.get(Document, audit.document_id) if audit.document_id else None
        audit_name = document.original_filename.replace(Path(document.original_filename).suffix, "") if document else f"audit_{audit.external_id}"
        date_str = datetime.utcnow().strftime("%Y-%m-%d")
        filename = f"final_report_{audit_name}_{date_str}.docx"
        
        return send_file(
            docx_buffer,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logger.exception(f"Error generating DOCX report for audit {audit_id}: {e}")
        return jsonify({"error": str(e)}), 500


def _final_report_to_markdown(report, audit) -> str:
    """Convert FinalReport to markdown format."""
    lines = [
        f"# Final Compliance Report: {audit.external_id}",
        "",
        "## Executive Summary",
        "",
        report.executive_summary,
        "",
    ]
    
    if report.critical_issues:
        lines.extend([
            "## Critical Issues",
            "",
        ])
        for issue in report.critical_issues:
            lines.extend([
                f"### {issue.get('title', 'Critical Issue')}",
                "",
                issue.get("description", ""),
                "",
            ])
            if issue.get("recommendations"):
                lines.append("**Recommendations:**")
                for rec in issue["recommendations"]:
                    lines.append(f"- {rec}")
                lines.append("")
    
    if report.warnings:
        lines.extend([
            "## Warnings",
            "",
        ])
        for warning in report.warnings:
            lines.extend([
                f"### {warning.get('title', 'Warning')}",
                "",
                warning.get("description", ""),
                "",
            ])
            if warning.get("recommendations"):
                lines.append("**Recommendations:**")
                for rec in warning["recommendations"]:
                    lines.append(f"- {rec}")
                lines.append("")
    
    if report.recommendations:
        lines.extend([
            "## Overall Recommendations",
            "",
        ])
        for rec in report.recommendations:
            lines.append(f"- {rec}")
        lines.append("")
    
    lines.extend([
        "## Overall Assessment",
        "",
        report.overall_assessment,
        "",
    ])
    
    return "\n".join(lines)

